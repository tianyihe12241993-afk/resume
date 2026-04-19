#!/usr/bin/env python3
"""Tailor a base resume (.docx) to each job URL in jobs.txt.

Outputs per job:
  output/<Company>_<Role>.docx
  output/<Company>_<Role>.pdf
  output/applications.csv  (job_url, company, role, location, docx, pdf)
"""

from __future__ import annotations

import argparse
import csv
import json
import os
import re
import sys
import time
import traceback
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional
from urllib.parse import urlparse, parse_qs

import requests
from bs4 import BeautifulSoup
from docx import Document
from dotenv import load_dotenv
from anthropic import Anthropic

load_dotenv()

ROOT = Path(__file__).parent
DEFAULT_RESUME = ROOT / "Tianyi.docx"
DEFAULT_JOBS = ROOT / "jobs.txt"
DEFAULT_OUT = ROOT / "output"

TAILOR_MODEL = os.getenv("TAILOR_MODEL", "claude-sonnet-4-6")
EXTRACT_MODEL = os.getenv("EXTRACT_MODEL", "claude-haiku-4-5-20251001")

UA = (
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"
)
HEADERS = {"User-Agent": UA, "Accept": "text/html,application/json,*/*"}


# --------------------------------------------------------------------------
# Resume parsing
# --------------------------------------------------------------------------

@dataclass
class JobBlock:
    title_idx: int
    company_idx: int
    title_text: str           # "Senior Software ENGINEER \tJul 2025 – Present"
    company_text: str         # "Roblox\tSan Mateo, CA, United States"
    bullet_indices: list = field(default_factory=list)
    bullets: list = field(default_factory=list)


@dataclass
class SkillBlock:
    category_idx: int
    items_idx: int
    category: str
    items: str


@dataclass
class ResumeStruct:
    summary_idx: int
    summary: str
    jobs: list
    skills: list


def _style_name(p) -> str:
    try:
        return p.style.name or ""
    except Exception:
        return ""


def parse_resume(doc: Document) -> ResumeStruct:
    paras = doc.paragraphs
    summary_idx: Optional[int] = None
    summary_text = ""
    jobs: list = []
    skills: list = []

    section = None
    current_job: Optional[JobBlock] = None
    pending_skill_cat: Optional[tuple] = None

    for i, p in enumerate(paras):
        style = _style_name(p)
        raw = p.text
        text = raw.strip()
        if not text:
            continue

        if style == "Heading 1":
            low = text.lower()
            if "summary" in low:
                section = "summary"
            elif "experience" in low:
                section = "experience"
            elif "skill" in low:
                section = "skills"
            elif "education" in low:
                section = "education"
            else:
                section = None
            continue

        if section == "summary" and style == "Normal" and summary_idx is None:
            summary_idx = i
            summary_text = raw

        elif section == "experience":
            if style == "Heading 2":
                # Title/date lines contain a 4-digit year, company/location lines don't.
                if re.search(r"\b(19|20)\d{2}\b", text) or "Present" in text:
                    current_job = JobBlock(
                        title_idx=i, company_idx=-1,
                        title_text=raw, company_text="",
                    )
                    jobs.append(current_job)
                elif current_job and current_job.company_idx == -1:
                    current_job.company_idx = i
                    current_job.company_text = raw
            elif style == "List Bullet" and current_job:
                current_job.bullet_indices.append(i)
                current_job.bullets.append(raw)

        elif section == "skills":
            if style == "List Paragraph":
                pending_skill_cat = (i, raw)
            elif style == "Normal" and pending_skill_cat:
                cat_i, cat_text = pending_skill_cat
                skills.append(SkillBlock(
                    category_idx=cat_i, items_idx=i,
                    category=cat_text, items=raw,
                ))
                pending_skill_cat = None

    if summary_idx is None:
        raise RuntimeError("Could not locate a Summary section in the resume.")

    return ResumeStruct(
        summary_idx=summary_idx, summary=summary_text,
        jobs=jobs, skills=skills,
    )


def set_paragraph_text(p, new_text: str) -> None:
    """Replace paragraph text while keeping paragraph style + first-run formatting."""
    runs = list(p.runs)
    if not runs:
        p.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


# --------------------------------------------------------------------------
# Job URL fetching (per-board fallbacks)
# --------------------------------------------------------------------------

def _get(url: str, timeout: int = 20) -> requests.Response:
    return requests.get(url, headers=HEADERS, timeout=timeout, allow_redirects=True)


def _fetch_ashby(url: str) -> Optional[dict]:
    """Ashby provides a public posting API per org."""
    # Extract slug + job id from /ashbyhq/<slug>/<jobid>[...]
    path = urlparse(url).path.strip("/").split("/")
    if len(path) < 2:
        return None
    slug, job_id = path[0], path[1]
    api = f"https://api.ashbyhq.com/posting-api/job-board/{slug}?includeCompensation=true"
    try:
        r = _get(api)
        if r.status_code != 200:
            return None
        data = r.json()
        for posting in data.get("jobs", []):
            if posting.get("id") == job_id:
                return {
                    "company": data.get("name") or slug,
                    "title": posting.get("title", ""),
                    "location": posting.get("location", ""),
                    "description": posting.get("descriptionPlain")
                        or BeautifulSoup(posting.get("descriptionHtml", ""),
                                         "html.parser").get_text("\n"),
                }
    except Exception:
        return None
    return None


def _fetch_lever(url: str) -> Optional[dict]:
    """Lever public posting API."""
    path = urlparse(url).path.strip("/").split("/")
    if len(path) < 2:
        return None
    slug, job_id = path[0], path[1]
    api = f"https://api.lever.co/v0/postings/{slug}/{job_id}?mode=json"
    try:
        r = _get(api)
        if r.status_code != 200:
            return None
        data = r.json()
        desc_html = data.get("description", "") + "\n"
        for block in data.get("lists", []):
            desc_html += f"\n<h3>{block.get('text','')}</h3>" + block.get("content", "")
        desc_html += "\n" + data.get("additional", "")
        return {
            "company": slug,
            "title": data.get("text", ""),
            "location": (data.get("categories") or {}).get("location", ""),
            "description": BeautifulSoup(desc_html, "html.parser").get_text("\n"),
        }
    except Exception:
        return None


def _fetch_greenhouse(url: str) -> Optional[dict]:
    """Greenhouse embed URLs → try the board API."""
    qs = parse_qs(urlparse(url).query)
    slug = (qs.get("for") or [None])[0]
    token = (qs.get("token") or [None])[0]
    if not slug or not token:
        return None
    api = f"https://boards-api.greenhouse.io/v1/boards/{slug}/jobs/{token}"
    try:
        r = _get(api)
        if r.status_code == 200:
            data = r.json()
            return {
                "company": (data.get("company_name")
                            or slug.replace("_", " ").title()),
                "title": data.get("title", ""),
                "location": (data.get("location") or {}).get("name", ""),
                "description": BeautifulSoup(data.get("content", ""),
                                             "html.parser").get_text("\n"),
            }
    except Exception:
        pass
    # Fallback: scrape the embed HTML
    try:
        r = _get(url)
        if r.status_code == 200:
            return _extract_from_html(r.text, fallback_company=slug)
    except Exception:
        pass
    return None


def _extract_from_html(html: str, fallback_company: str = "") -> dict:
    soup = BeautifulSoup(html, "html.parser")
    title = ""
    for sel in ["h1", ".app-title", "[data-test='job-title']"]:
        el = soup.select_one(sel)
        if el and el.get_text(strip=True):
            title = el.get_text(strip=True)
            break
    text = soup.get_text("\n")
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return {
        "company": fallback_company,
        "title": title,
        "location": "",
        "description": text,
    }


def fetch_job_posting(url: str) -> dict:
    """Return {company, title, location, description}.

    Tries board-specific APIs first, then falls back to generic HTML scrape.
    """
    host = urlparse(url).netloc.lower()

    # Board-specific paths
    if "ashbyhq.com" in host:
        info = _fetch_ashby(url)
        if info and info.get("description"):
            return info
    if "lever.co" in host:
        info = _fetch_lever(url)
        if info and info.get("description"):
            return info
    if "greenhouse.io" in host:
        info = _fetch_greenhouse(url)
        if info and info.get("description"):
            return info

    # Generic fetch
    candidates = [url]
    if url.endswith("/application"):
        candidates.append(url.rsplit("/application", 1)[0])
    if url.endswith("/apply"):
        candidates.append(url.rsplit("/apply", 1)[0])

    last_err = None
    for u in candidates:
        try:
            r = _get(u)
            if r.status_code == 200 and len(r.text) > 500:
                return _extract_from_html(r.text)
        except Exception as e:
            last_err = e
    raise RuntimeError(f"Could not fetch job page: {url} ({last_err})")


# --------------------------------------------------------------------------
# Claude: enrich job info + tailor resume
# --------------------------------------------------------------------------

def _client() -> Anthropic:
    key = os.getenv("ANTHROPIC_API_KEY")
    if not key:
        print("ERROR: ANTHROPIC_API_KEY is not set. Copy .env.sample to .env "
              "and fill in your key.", file=sys.stderr)
        sys.exit(1)
    return Anthropic(api_key=key)


def _extract_tagged_json(text: str) -> dict:
    """Pull JSON out of a Claude response that may wrap it in <json> tags or markdown."""
    m = re.search(r"<json>\s*(\{.*?\})\s*</json>", text, re.DOTALL)
    if m:
        return json.loads(m.group(1))
    m = re.search(r"```json\s*(\{.*?\})\s*```", text, re.DOTALL)
    if m:
        return json.loads(m.group(1))
    m = re.search(r"(\{.*\})", text, re.DOTALL)
    if m:
        return json.loads(m.group(1))
    raise ValueError(f"No JSON in response: {text[:400]}")


def normalize_job_info(client: Anthropic, raw: dict, url: str) -> dict:
    """If company/title/description look thin, ask Claude to clean them up."""
    desc = (raw.get("description") or "").strip()
    if len(desc) >= 400 and raw.get("company") and raw.get("title"):
        return raw

    prompt = f"""Clean up and normalize this job posting metadata. The raw extraction may contain navigation text, boilerplate, or truncation. Preserve all actual job-description content.

URL: {url}

RAW:
{json.dumps(raw, ensure_ascii=False)[:12000]}

Return <json>{{"company": "...", "title": "...", "location": "...", "description": "..."}}</json>."""

    resp = client.messages.create(
        model=EXTRACT_MODEL,
        max_tokens=2000,
        messages=[{"role": "user", "content": prompt}],
    )
    txt = resp.content[0].text
    cleaned = _extract_tagged_json(txt)
    return {
        "company": cleaned.get("company") or raw.get("company", ""),
        "title": cleaned.get("title") or raw.get("title", ""),
        "location": cleaned.get("location") or raw.get("location", ""),
        "description": cleaned.get("description") or desc,
    }


TAILOR_SYSTEM = """You are an expert resume editor. You will tailor a candidate's resume to a specific job posting without fabricating anything.

HARD RULES — violating any of these invalidates the output:
1. NEVER invent technologies, companies, projects, metrics, titles, dates, or responsibilities that are not already present in the candidate's resume. Only rewording and reordering are allowed.
2. Preserve the EXACT number of bullets for each job. Do not add or drop bullets.
3. Preserve the EXACT number of skill categories and do not drop any of them. You may reorder items inside a category and reword a category label slightly, but do not add items that are not already listed for the candidate.
4. Every bullet must stay TRUE to the original fact: same employer, same tech stack, same outcome, same order of magnitude for metrics. Surface-level rewording only.
5. Keep the candidate's existing action-verb / third-person style. Do not use "I" or "we".

TAILORING OBJECTIVES:
- Rewrite the Summary (3–5 sentences) to front-load the experience most relevant to the target role. Keep total length close to the original.
- Reorder bullets within each job so the most JD-relevant bullet comes first.
- Subtly reword bullets to surface JD keywords WHEN TRUTHFUL (e.g., if the JD emphasizes "distributed systems" and a bullet says "microservices infrastructure", either phrasing is fine; if the JD is about iOS and the bullet is about backend, leave it alone).
- Reorder skill categories so the most JD-relevant ones come first. Reorder items inside each category the same way.
- Do not mention the target company or job title inside the resume body.

OUTPUT FORMAT — you MUST call the `submit_tailored_resume` tool exactly once with the tailored content. Do not emit text, markdown, or JSON outside the tool call.

The jobs array MUST be in the same order as the input resume, with the same number of bullets per job.
The skills array MUST contain the same categories as the input (in your chosen order), each with its items as a single comma-separated string.
Do not include apostrophes formatted as typographic quotes (use straight ASCII apostrophes instead) if it risks formatting issues."""


SUBMIT_TOOL = {
    "name": "submit_tailored_resume",
    "description": "Submit the tailored resume content. Call exactly once.",
    "input_schema": {
        "type": "object",
        "required": ["summary", "jobs", "skills"],
        "properties": {
            "summary": {
                "type": "string",
                "description": "Tailored 3–5 sentence professional summary.",
            },
            "jobs": {
                "type": "array",
                "description": "Tailored jobs, in the SAME order as the input, with the SAME number of bullets each.",
                "items": {
                    "type": "object",
                    "required": ["title_hint", "bullets"],
                    "properties": {
                        "title_hint": {"type": "string"},
                        "bullets": {
                            "type": "array",
                            "items": {"type": "string"},
                        },
                    },
                },
            },
            "skills": {
                "type": "array",
                "description": "Skill categories in tailored order. Must contain the same set of categories as the input.",
                "items": {
                    "type": "object",
                    "required": ["category", "items"],
                    "properties": {
                        "category": {"type": "string"},
                        "items": {"type": "string"},
                    },
                },
            },
        },
    },
}


def tailor_resume(client: Anthropic, resume: ResumeStruct, job: dict) -> dict:
    """Call Claude to produce tailored summary/bullets/skills."""
    resume_json = {
        "summary": resume.summary.strip(),
        "jobs": [
            {"title_hint": j.title_text.strip(),
             "bullets": [b.strip() for b in j.bullets]}
            for j in resume.jobs
        ],
        "skills": [
            {"category": s.category.strip(), "items": s.items.strip()}
            for s in resume.skills
        ],
    }

    user_msg = [
        {
            "type": "text",
            "text": (
                "Candidate resume (editable portions):\n<resume>\n"
                + json.dumps(resume_json, ensure_ascii=False, indent=2)
                + "\n</resume>"
            ),
            "cache_control": {"type": "ephemeral"},
        },
        {
            "type": "text",
            "text": (
                "Target job posting:\n<job>\n"
                f"Company: {job.get('company','')}\n"
                f"Title: {job.get('title','')}\n"
                f"Location: {job.get('location','')}\n\n"
                f"Description:\n{(job.get('description') or '')[:15000]}\n"
                "</job>\n\n"
                "Produce the tailored resume now, following every HARD RULE."
            ),
        },
    ]

    resp = client.messages.create(
        model=TAILOR_MODEL,
        max_tokens=4000,
        system=TAILOR_SYSTEM,
        tools=[SUBMIT_TOOL],
        tool_choice={"type": "tool", "name": "submit_tailored_resume"},
        messages=[{"role": "user", "content": user_msg}],
    )
    out = None
    for block in resp.content:
        if getattr(block, "type", None) == "tool_use" and block.name == "submit_tailored_resume":
            out = block.input
            break
    if out is None:
        raise RuntimeError(
            f"Model did not call submit_tailored_resume. stop_reason={resp.stop_reason}"
        )
    _validate_tailored(out, resume)
    return out


def _validate_tailored(out: dict, resume: ResumeStruct) -> None:
    if not isinstance(out.get("summary"), str) or not out["summary"].strip():
        raise ValueError("Tailored output missing summary.")
    jobs = out.get("jobs") or []
    if len(jobs) != len(resume.jobs):
        raise ValueError(f"Jobs count mismatch: got {len(jobs)}, expected {len(resume.jobs)}.")
    for i, (tj, rj) in enumerate(zip(jobs, resume.jobs)):
        if len(tj.get("bullets") or []) != len(rj.bullets):
            raise ValueError(
                f"Job {i} bullet count mismatch: got {len(tj.get('bullets') or [])}, "
                f"expected {len(rj.bullets)}."
            )
    skills = out.get("skills") or []
    if len(skills) != len(resume.skills):
        raise ValueError(
            f"Skills category count mismatch: got {len(skills)}, expected {len(resume.skills)}."
        )


# --------------------------------------------------------------------------
# Apply tailored output back to a fresh copy of the docx
# --------------------------------------------------------------------------

def apply_tailoring(src_docx: Path, resume: ResumeStruct,
                    tailored: dict, dst_docx: Path) -> None:
    doc = Document(str(src_docx))
    paras = doc.paragraphs

    # Summary
    set_paragraph_text(paras[resume.summary_idx], tailored["summary"].strip())

    # Bullets — positions stay the same; we never change titles/dates/companies.
    for job, tjob in zip(resume.jobs, tailored["jobs"]):
        for idx, new_bullet in zip(job.bullet_indices, tjob["bullets"]):
            set_paragraph_text(paras[idx], new_bullet.strip())

    # Skills — we may reorder categories, but paragraph indices are fixed.
    # Simplest: write tailored skills back in the SAME positions as originals,
    # mapping by order (since model is told to keep the schema). That means
    # "reordered" categories land in the order the model returned them.
    for skill, tskill in zip(resume.skills, tailored["skills"]):
        set_paragraph_text(paras[skill.category_idx], tskill["category"].strip())
        set_paragraph_text(paras[skill.items_idx], tskill["items"].strip())

    dst_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst_docx))


# --------------------------------------------------------------------------
# PDF conversion
# --------------------------------------------------------------------------

def convert_to_pdf(docx_path: Path) -> Optional[Path]:
    pdf_path = docx_path.with_suffix(".pdf")
    # Try docx2pdf first (uses MS Word on macOS/Windows, LibreOffice on Linux).
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        if pdf_path.exists():
            return pdf_path
    except Exception as e:
        print(f"  ⚠ docx2pdf failed: {e}", file=sys.stderr)

    # Fallback: LibreOffice headless if installed.
    import shutil, subprocess
    soffice = shutil.which("soffice") or shutil.which("libreoffice") \
        or "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    if Path(soffice).exists():
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf",
                 "--outdir", str(pdf_path.parent), str(docx_path)],
                check=True, capture_output=True,
            )
            if pdf_path.exists():
                return pdf_path
        except Exception as e:
            print(f"  ⚠ LibreOffice conversion failed: {e}", file=sys.stderr)
    return None


# --------------------------------------------------------------------------
# Utilities
# --------------------------------------------------------------------------

def read_jobs_file(path: Path) -> list:
    urls = []
    for line in path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        urls.append(line)
    return urls


def safe_slug(s: str, limit: int = 60) -> str:
    s = (s or "Unknown").strip()
    s = re.sub(r"[^A-Za-z0-9._ -]+", "", s)
    s = re.sub(r"\s+", "_", s)
    return s[:limit] or "Unknown"


# --------------------------------------------------------------------------
# Main
# --------------------------------------------------------------------------

def main() -> int:
    ap = argparse.ArgumentParser(description="Tailor a resume to many job URLs.")
    ap.add_argument("--resume", type=Path, default=DEFAULT_RESUME,
                    help=f"Path to base .docx (default: {DEFAULT_RESUME.name})")
    ap.add_argument("--jobs", type=Path, default=DEFAULT_JOBS,
                    help=f"Path to jobs list (default: {DEFAULT_JOBS.name})")
    ap.add_argument("--out", type=Path, default=DEFAULT_OUT,
                    help="Output directory (default: output/)")
    ap.add_argument("--limit", type=int, default=0,
                    help="Stop after N jobs (0 = all)")
    ap.add_argument("--skip-pdf", action="store_true",
                    help="Do not attempt PDF conversion")
    args = ap.parse_args()

    if not args.resume.exists():
        print(f"ERROR: resume not found: {args.resume}", file=sys.stderr)
        return 1
    if not args.jobs.exists():
        print(f"ERROR: jobs file not found: {args.jobs}", file=sys.stderr)
        return 1

    urls = read_jobs_file(args.jobs)
    if args.limit:
        urls = urls[: args.limit]
    if not urls:
        print("No URLs to process.", file=sys.stderr)
        return 1

    args.out.mkdir(parents=True, exist_ok=True)
    client = _client()

    print(f"Parsing base resume: {args.resume.name}")
    base_doc = Document(str(args.resume))
    resume = parse_resume(base_doc)
    print(f"  summary@{resume.summary_idx}, "
          f"{len(resume.jobs)} jobs, "
          f"{sum(len(j.bullets) for j in resume.jobs)} bullets, "
          f"{len(resume.skills)} skill categories")

    csv_path = args.out / "applications.csv"
    write_header = not csv_path.exists()
    csv_f = csv_path.open("a", newline="", encoding="utf-8")
    writer = csv.writer(csv_f)
    if write_header:
        writer.writerow(["job_url", "company", "role", "location", "docx", "pdf", "status"])

    for i, url in enumerate(urls, 1):
        print(f"\n[{i}/{len(urls)}] {url}")
        status, company, role, location, docx_rel, pdf_rel = "ok", "", "", "", "", ""
        try:
            raw = fetch_job_posting(url)
            info = normalize_job_info(client, raw, url)
            company = info.get("company", "").strip() or "Unknown"
            role = info.get("title", "").strip() or "Role"
            location = info.get("location", "").strip()
            print(f"  → {company} — {role} ({location or 'n/a'})")

            tailored = tailor_resume(client, resume, info)
            slug = f"{safe_slug(company)}_{safe_slug(role)}"
            docx_path = args.out / f"{slug}.docx"
            apply_tailoring(args.resume, resume, tailored, docx_path)
            print(f"  ✓ docx → {docx_path.name}")
            docx_rel = docx_path.name

            if not args.skip_pdf:
                pdf_path = convert_to_pdf(docx_path)
                if pdf_path and pdf_path.exists():
                    print(f"  ✓ pdf  → {pdf_path.name}")
                    pdf_rel = pdf_path.name
                else:
                    status = "ok_no_pdf"
                    print("  ⚠ skipped PDF (no converter worked)")
        except Exception as e:
            status = f"error: {type(e).__name__}"
            print(f"  ✗ {e}", file=sys.stderr)
            traceback.print_exc(limit=2)

        writer.writerow([url, company, role, location, docx_rel, pdf_rel, status])
        csv_f.flush()

    csv_f.close()
    print(f"\nDone. CSV → {csv_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
