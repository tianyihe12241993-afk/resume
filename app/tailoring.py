"""Resume parsing + Claude-powered tailoring + docx writing."""
from __future__ import annotations

import copy as _copy
import hashlib
import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from anthropic import Anthropic
from docx import Document
from docx.oxml.ns import qn

from . import config

# --------------------------------------------------------------------------
# Resume structure
# --------------------------------------------------------------------------

@dataclass
class JobBlock:
    title_idx: int
    company_idx: int
    title_text: str
    company_text: str
    bullet_indices: list = field(default_factory=list)
    bullets: list = field(default_factory=list)


@dataclass
class SkillBlock:
    category_idx: int
    items_idx: int
    category: str
    items: str


@dataclass
class ResumeStruct:
    summary_idx: int           # -1 if no summary block found
    summary: str
    jobs: list
    skills: list

    @property
    def has_summary(self) -> bool:
        return self.summary_idx >= 0


def _style_name(p) -> str:
    try:
        return p.style.name or ""
    except Exception:
        return ""


# --------------------------------------------------------------------------
# AI-based resume structure analysis (adapts to any base-resume layout)
# --------------------------------------------------------------------------

_STRUCTURE_SYSTEM = """You analyze the structure of a resume document to identify which paragraphs belong to which section. You are given a numbered list of non-empty paragraphs; return a JSON document that maps indices to sections.

Output STRICT JSON only, with this exact shape:

{
  "summary_indices": [<paragraph indices making up the Summary / Profile / Objective block, in order>],
  "jobs": [
    {
      "title_idx": <index of the paragraph containing the role/company line>,
      "bullet_indices": [<indices of bullet points under this job, in order>]
    },
    ...
  ],
  "skills": [
    {
      "category_idx": <index of the skill category label>,
      "items_idx": <index of the paragraph listing items in that category>
    },
    ...
  ]
}

Rules:
- Every index you return MUST appear in the input.
- If a section doesn't exist in the resume, return an empty list for it.
- NEVER include section-header paragraphs (like "EXPERIENCE", "SKILLS", "SUMMARY") in any of the output arrays.
- For jobs: `title_idx` points to the paragraph identifying the role/company (not date lines, not location lines, not sub-headers).
- For bullets: include only actual responsibility/achievement bullets, not title/company/date lines.
- For skills: if a skill category and its items share a single paragraph (e.g. "Languages: Python, Go"), set category_idx == items_idx and the tailoring step will treat the whole line as rewrite-target.
- Never invent indices; only use indices present in the input.
- Omit education, certifications, awards, hobbies — they are not tailored.
"""


def _analyze_structure_with_claude(items: list[tuple[int, str, str]]) -> dict:
    """Send paragraphs to Claude-Haiku and get structured section indices back."""
    payload = [{"i": i, "style": s, "text": t[:280]} for i, s, t in items]
    user = "Paragraphs:\n" + json.dumps(payload, ensure_ascii=False)

    client = _client()
    resp = client.messages.create(
        model=config.EXTRACT_MODEL,
        max_tokens=3000,
        temperature=0,
        system=_STRUCTURE_SYSTEM,
        messages=[{"role": "user", "content": user}],
    )
    text = "".join(b.text for b in resp.content if getattr(b, "type", None) == "text")
    m = re.search(r"\{.*\}", text, re.DOTALL)
    if not m:
        raise RuntimeError(f"Structure analyzer returned no JSON: {text[:400]}")
    return json.loads(m.group(0))


# In-memory cache: { sha256(file) -> structure dict }. Avoids re-calling
# Claude for every URL in a batch since the base resume doesn't change.
_structure_cache: dict[str, dict] = {}


def _docx_sha256(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        while True:
            chunk = f.read(65536)
            if not chunk:
                break
            h.update(chunk)
    return h.hexdigest()


def parse_resume_from_path(path: Path) -> ResumeStruct:
    """Parse the base resume, using AI to detect layout. Cached by file hash."""
    digest = _docx_sha256(path)
    doc = Document(str(path))
    paras = doc.paragraphs

    # Collect non-empty paragraphs with style info
    triples: list[tuple[int, str, str]] = []
    for i, p in enumerate(paras):
        text = p.text.strip()
        if text:
            triples.append((i, _style_name(p), text))
    if not triples:
        raise RuntimeError("The base resume has no text.")

    analysis = _structure_cache.get(digest)
    if analysis is None:
        try:
            analysis = _analyze_structure_with_claude(triples)
            _structure_cache[digest] = analysis
        except Exception:
            # Fallback to the legacy heuristic parser (Heading 1/Normal styles).
            analysis = _heuristic_analyze(triples)
            _structure_cache[digest] = analysis

    return _build_struct_from_analysis(paras, analysis)


def parse_resume(doc: Document) -> ResumeStruct:
    """Back-compat entry point — no cache, AI analyzes every call."""
    paras = doc.paragraphs
    triples: list[tuple[int, str, str]] = []
    for i, p in enumerate(paras):
        text = p.text.strip()
        if text:
            triples.append((i, _style_name(p), text))
    if not triples:
        raise RuntimeError("The base resume has no text.")
    try:
        analysis = _analyze_structure_with_claude(triples)
    except Exception:
        analysis = _heuristic_analyze(triples)
    return _build_struct_from_analysis(paras, analysis)


def _build_struct_from_analysis(paras: list, analysis: dict) -> ResumeStruct:
    valid_idx = {i for i in range(len(paras))}

    summary_indices = [i for i in (analysis.get("summary_indices") or []) if i in valid_idx]
    summary_idx = summary_indices[0] if summary_indices else -1
    summary_text = paras[summary_idx].text if summary_idx >= 0 else ""

    jobs: list = []
    for ja in (analysis.get("jobs") or []):
        t = ja.get("title_idx")
        b = [i for i in (ja.get("bullet_indices") or []) if i in valid_idx]
        if t is None or t not in valid_idx:
            continue
        jobs.append(JobBlock(
            title_idx=t, company_idx=-1,
            title_text=paras[t].text, company_text="",
            bullet_indices=b,
            bullets=[paras[i].text for i in b],
        ))

    skills: list = []
    for sa in (analysis.get("skills") or []):
        c = sa.get("category_idx")
        it = sa.get("items_idx")
        if c is None or it is None or c not in valid_idx or it not in valid_idx:
            continue
        skills.append(SkillBlock(
            category_idx=c, items_idx=it,
            category=paras[c].text, items=paras[it].text,
        ))

    if summary_idx < 0 and not jobs and not skills:
        raise RuntimeError(
            "Could not identify any resume sections to tailor. "
            "Make sure the resume has at least a summary, experience bullets, or skills."
        )
    return ResumeStruct(
        summary_idx=summary_idx, summary=summary_text,
        jobs=jobs, skills=skills,
    )


def _heuristic_analyze(items: list[tuple[int, str, str]]) -> dict:
    """Fallback when the AI analyzer is unavailable: style-based heuristic."""
    section = None
    summary_indices: list[int] = []
    jobs: list = []
    cur_job: Optional[dict] = None
    skills: list = []
    pending_cat: Optional[tuple[int, str]] = None

    for i, style, text in items:
        low = text.lower().strip(":").strip()
        if style == "Heading 1" or low in ("summary", "profile", "objective", "professional summary",
                                             "experience", "work experience", "professional experience",
                                             "skills", "technical skills", "core skills",
                                             "education"):
            if "summary" in low or "profile" in low or "objective" in low:
                section = "summary"
            elif "experience" in low:
                section = "experience"
                cur_job = None
            elif "skill" in low:
                section = "skills"
            elif "education" in low:
                section = "education"
            else:
                section = None
            continue

        if section == "summary":
            summary_indices.append(i)
        elif section == "experience":
            if style == "Heading 2":
                if re.search(r"\b(19|20)\d{2}\b", text) or "Present" in text:
                    cur_job = {"title_idx": i, "bullet_indices": []}
                    jobs.append(cur_job)
                elif cur_job is None:
                    cur_job = {"title_idx": i, "bullet_indices": []}
                    jobs.append(cur_job)
            elif style == "List Bullet" and cur_job:
                cur_job["bullet_indices"].append(i)
            elif text.lstrip().startswith(("•", "-", "·", "*")) and cur_job:
                cur_job["bullet_indices"].append(i)
        elif section == "skills":
            if style == "List Paragraph":
                pending_cat = (i, text)
            elif pending_cat is not None:
                skills.append({"category_idx": pending_cat[0], "items_idx": i})
                pending_cat = None
            elif ":" in text:
                skills.append({"category_idx": i, "items_idx": i})

    return {"summary_indices": summary_indices, "jobs": jobs, "skills": skills}


_LEAD_TERMINATORS = (":", "—", "–", " - ", " · ", "•")


def _split_lead_in_new_text(new_text: str) -> Optional[int]:
    """Return the index AFTER the lead-in terminator (':', '—', etc) in new_text,
    if there is one near the start of the bullet. None if the new text doesn't
    fit the labeled-bullet pattern."""
    # Only consider the first ~80 chars — beyond that it's not a lead-in.
    head = new_text[:80]
    candidates: list[int] = []
    for term in _LEAD_TERMINATORS:
        idx = head.find(term)
        if idx > 0:
            candidates.append(idx + len(term))
    if not candidates:
        return None
    # Use the earliest terminator.
    return min(candidates)


def set_paragraph_text(p, new_text: str) -> None:
    """Replace a paragraph's text, preserving its formatting structure.

    Three patterns the resume parser sees regularly:
      1. A single run (uniform formatting) → easy: just replace text.
      2. "**Bold lead-in:** plain body" — the labeled-bullet pattern. Preserve
         the bold/plain split if the new text also has a lead terminator.
      3. Anything else → keep whichever run carried the most characters
         (majority-wins) so the dominant formatting is preserved.
    """
    runs = list(p.runs)
    if not runs:
        p.add_run(new_text)
        return

    if len(runs) == 1:
        runs[0].text = new_text
        return

    # Try to detect the labeled-bullet pattern: a bold first run that ends
    # with a label terminator, followed by a non-bold run carrying the body.
    first = runs[0]
    first_text = (first.text or "").rstrip()
    looks_like_label = (
        bool(first.bold)
        and any(first_text.endswith(t.strip()) for t in _LEAD_TERMINATORS if t.strip())
        and any((not r.bold) and (r.text or "").strip() for r in runs[1:])
    )
    if looks_like_label:
        split_at = _split_lead_in_new_text(new_text)
        if split_at is not None:
            lead = new_text[:split_at].rstrip()
            body = new_text[split_at:]
            # Preserve a leading space so "Lead: body" doesn't collapse to "Lead:body".
            if body and not body.startswith(" "):
                body = " " + body
            first.text = lead
            plain_run = next((r for r in runs[1:] if not r.bold), runs[1])
            plain_run.text = body
            for r in runs:
                if r is not first and r is not plain_run:
                    r._element.getparent().remove(r._element)
            return

    # Uniformly-bold paragraph → keep it bold.
    text_runs = [r for r in runs if (r.text or "").strip()]
    if text_runs and all(bool(r.bold) for r in text_runs):
        keep = max(text_runs, key=lambda r: len(r.text or ""))
        keep.bold = True
        keep.text = new_text
        for r in runs:
            if r is not keep:
                r._element.getparent().remove(r._element)
        return

    # Mixed emphasis (some runs bold, some plain). Use the dominant plain run as
    # the base formatting, then re-bold any originally-bold word that survived
    # verbatim into the rewrite — so emphasis on terms the rewriter kept (e.g. a
    # bolded "Kubernetes") is restored rather than flattened.
    bold_subs = [
        (r.text or "").strip()
        for r in runs
        if bool(r.bold) and (r.text or "").strip()
    ]
    plain_candidates = [r for r in runs if not bool(r.bold)]
    keep = max(plain_candidates or runs, key=lambda r: len(r.text or ""))
    for r in runs:
        if r is not keep:
            r._element.getparent().remove(r._element)
    keep.bold = False

    segments = _bold_segments(new_text, bold_subs)
    if not segments:
        keep.text = new_text
        return
    _rebuild_run_with_bold(p, keep, segments)


def _bold_segments(new_text: str, bold_subs: list[str]) -> Optional[list[tuple[str, bool]]]:
    """Split `new_text` into (text, is_bold) segments by marking every verbatim
    occurrence of an originally-bold substring as bold. Returns None when no
    bold substring survives (caller then writes plain text)."""
    n = len(new_text)
    if not n:
        return None
    mask = [False] * n
    # Longer substrings first so a phrase wins over an incidental short word.
    subs = sorted(
        {s for s in bold_subs if len(s) >= 2 and any(ch.isalnum() for ch in s)},
        key=len, reverse=True,
    )
    any_bold = False
    for s in subs:
        start = 0
        while True:
            i = new_text.find(s, start)
            if i < 0:
                break
            for k in range(i, i + len(s)):
                mask[k] = True
            any_bold = True
            start = i + len(s)
    if not any_bold:
        return None

    segments: list[tuple[str, bool]] = []
    cur = mask[0]
    buf = new_text[0]
    for idx in range(1, n):
        if mask[idx] == cur:
            buf += new_text[idx]
        else:
            segments.append((buf, cur))
            buf = new_text[idx]
            cur = mask[idx]
    segments.append((buf, cur))
    return segments


def _set_run_rpr(run, rpr_clone) -> None:
    """Replace a run's properties with a clone of another run's rPr (font, size,
    color, etc.), so a freshly-added run inherits the source run's look."""
    el = run._element
    existing = el.find(qn("w:rPr"))
    if existing is not None:
        el.remove(existing)
    if rpr_clone is not None:
        el.insert(0, _copy.deepcopy(rpr_clone))


def _split_skill(category: str, items: str) -> tuple[str, str]:
    """Return (label, body) for a skill row. Handles both shapes the pipeline
    produces: clean separate fields (tailor_skills) and the full "Label: body"
    line duplicated into both fields (deterministic fallback / single-paragraph
    parse)."""
    category = (category or "").strip()
    items = (items or "").strip()
    if items and category and items != category and not items.startswith(category):
        return category.rstrip(":").strip(), items
    full = items or category
    if ":" in full:
        label, _, body = full.partition(":")
        return label.strip(), body.strip()
    return full.strip(), ""


def _apply_skill_row(p, category: str, items: str) -> None:
    """Write a skill row into a single paragraph as a bold 'Label:' followed by
    plain items, cloning the original bold/plain run formatting so fonts, sizes
    and colors survive. Used when the resume keeps category + items in ONE
    paragraph (category_idx == items_idx); writing them with two separate
    set_paragraph_text calls would make the second clobber the first."""
    label, body = _split_skill(category, items)
    runs = list(p.runs)
    bold_rpr = plain_rpr = None
    for r in runs:
        if not (r.text or "").strip():
            continue
        if r.bold and bold_rpr is None:
            bold_rpr = r._element.find(qn("w:rPr"))
        elif not r.bold and plain_rpr is None:
            plain_rpr = r._element.find(qn("w:rPr"))
    if bold_rpr is None:
        bold_rpr = plain_rpr
    if plain_rpr is None:
        plain_rpr = bold_rpr
    bold_clone = _copy.deepcopy(bold_rpr) if bold_rpr is not None else None
    plain_clone = _copy.deepcopy(plain_rpr) if plain_rpr is not None else None

    for r in runs:
        r._element.getparent().remove(r._element)

    label_run = p.add_run(label + (":" if body else ""))
    _set_run_rpr(label_run, bold_clone)
    label_run.bold = True
    if body:
        body_run = p.add_run(" " + body)
        _set_run_rpr(body_run, plain_clone)
        body_run.bold = False


def _rebuild_run_with_bold(p, keep, segments: list[tuple[str, bool]]) -> None:
    """Rewrite paragraph `p` as `segments`, cloning `keep`'s run properties
    (font, size, color) onto each new run so only bold differs between them."""
    rpr = keep._element.find(qn("w:rPr"))
    rpr_clone = _copy.deepcopy(rpr) if rpr is not None else None

    first_text, first_bold = segments[0]
    keep.text = first_text
    keep.bold = first_bold
    for text, is_bold in segments[1:]:
        r = p.add_run(text)  # keep is the only existing run, so this appends after it
        el = r._element
        existing = el.find(qn("w:rPr"))
        if existing is not None:
            el.remove(existing)
        if rpr_clone is not None:
            el.insert(0, _copy.deepcopy(rpr_clone))
        r.bold = is_bold


# --------------------------------------------------------------------------
# Claude calls
# --------------------------------------------------------------------------

def _client() -> Anthropic:
    if not config.ANTHROPIC_API_KEY:
        raise RuntimeError(
            "ANTHROPIC_API_KEY is not set. Edit .env and restart the server."
        )
    return Anthropic(api_key=config.ANTHROPIC_API_KEY)


def _extract_tagged_json(text: str) -> dict:
    m = re.search(r"<json>\s*(\{.*?\})\s*</json>", text, re.DOTALL)
    if m:
        return json.loads(m.group(1))
    m = re.search(r"```json\s*(\{.*?\})\s*```", text, re.DOTALL)
    if m:
        return json.loads(m.group(1))
    m = re.search(r"(\{.*\})", text, re.DOTALL)
    if m:
        return json.loads(m.group(1))
    raise ValueError(f"No JSON in response: {text[:400]}")


def normalize_job_info(raw: dict, url: str = "") -> dict:
    """Clean thin scrapes by asking Haiku to extract company/title/description."""
    desc = (raw.get("description") or "").strip()
    if len(desc) >= 400 and raw.get("company") and raw.get("title"):
        return raw

    client = _client()
    prompt = f"""Clean up and normalize this job posting metadata. The raw extraction may contain navigation text, boilerplate, or truncation. Preserve all actual job-description content.

URL: {url}

RAW:
{json.dumps(raw, ensure_ascii=False)[:12000]}

Return <json>{{"company": "...", "title": "...", "location": "...", "description": "..."}}</json>."""

    resp = client.messages.create(
        model=config.EXTRACT_MODEL,
        max_tokens=2000,
        temperature=0,
        messages=[{"role": "user", "content": prompt}],
    )
    txt = resp.content[0].text
    cleaned = _extract_tagged_json(txt)
    return {
        "company": cleaned.get("company") or raw.get("company", ""),
        "title": cleaned.get("title") or raw.get("title", ""),
        "location": cleaned.get("location") or raw.get("location", ""),
        "work_type": raw.get("work_type"),
        "description": cleaned.get("description") or desc,
    }


def extract_job_info_from_text(jd_text: str) -> dict:
    """Pull company / title / location from raw pasted JD text using Haiku.
    Used when the user manually pastes a JD instead of letting the scraper
    fetch it. Returns {"company": "", "title": "", "location": ""} on
    failure rather than raising — the manual flow is best-effort.
    """
    text = (jd_text or "").strip()
    if len(text) < 50:
        return {"company": "", "title": "", "location": ""}

    prompt = (
        "Extract structured metadata from this raw job description. The text "
        "may be a partial paste, may include boilerplate or company-about "
        "sections, and may not have an explicit \"Company:\" or \"Title:\" "
        "label. Use your judgment. If a field cannot be determined, use an "
        "empty string.\n\n"
        f"JD TEXT:\n{text[:12000]}\n\n"
        'Return <json>{"company": "...", "title": "...", "location": "..."}</json>.'
    )

    try:
        client = _client()
        resp = client.messages.create(
            model=config.EXTRACT_MODEL,
            max_tokens=400,
            temperature=0,
            messages=[{"role": "user", "content": prompt}],
        )
        txt = "".join(b.text for b in resp.content if getattr(b, "type", None) == "text")
        out = _extract_tagged_json(txt)
    except Exception:
        return {"company": "", "title": "", "location": ""}

    return {
        "company":  (out.get("company")  or "").strip(),
        "title":    (out.get("title")    or "").strip(),
        "location": (out.get("location") or "").strip(),
    }


TAILOR_SYSTEM = """You are an expert resume editor. Tailor the candidate's resume to a specific job posting by surfacing relevant existing experience as aggressively as the facts support — never by inventing anything.

WORKFLOW (think through internally before writing output):
1. Extract from the JD: required technologies, years of experience, domain/industry, role seniority, and the 5–7 top "must-have" keywords in the order the posting emphasizes them.
2. For each JD requirement, find the matching content in the resume. A match can be:
   - An exact match (same tech/domain word appears verbatim).
   - An adjacent match (the candidate did something equivalent — e.g. they built a "publisher dashboard" and the JD wants "content management").
   Prefer exact matches when ordering; use adjacent matches when exact ones don't exist.
3. If the JD has a hard requirement the resume doesn't support (e.g. "Strong Java" when Java appears only in skills), DO NOT fabricate. Surface whatever truthful Java-adjacent signal exists (e.g. the one bullet where Spring Boot was used) by moving it forward.
4. Rewrite the summary, reorder bullets, and reorder skill categories per the rules below.

HARD RULES — any violation invalidates the output:
1. NEVER invent job experience: no new bullets, no new employers, no fake projects, no metrics that did not exist, no titles or dates that aren't in the resume.
2. Preserve EXACTLY the same number of bullets per job. Same number of skill categories.
3. Preserve every metric (percentages, latencies, team sizes, user counts, revenue) verbatim within bullets. You may move them; you cannot change or remove them.
4. Preserve every tech-stack term, company name, and product name that appears in a bullet. Rewording the verbs and structure is allowed; the nouns that carry facts are NOT.
5. Third-person / action-verb voice throughout. No "I", "we", "my", "our".
6. Do NOT mention the target company or target job title inside the resume body.
7. SKILLS section is the ONE place you may add JD keywords the candidate doesn't already list (rules below). Bullets and summary stay grounded in the candidate's actual experience.

SUMMARY (3–5 sentences, ~60–110 words, matches input length ±20%):
- Sentence 1: years of experience + role identity + DOMAIN signal that matches the JD (pull the JD's domain word when possible, e.g. "fintech", "healthtech", "developer tools", "e-commerce").
- Sentence 2: 2–3 of the JD's top required technologies that the candidate actually has, with a real employer or project as proof ("shipped X at Walmart", not "worked with X").
- Sentence 3: most relevant current/recent work with scope (users / scale / latency / team).
- Optional sentence 4: leadership / cross-functional / team-size signal, if the JD values it and the candidate has it.
- Do NOT use filler ("passionate about...", "proven ability to..."). Every sentence must add a concrete fact.

BULLETS — reorder and reword:
- Reorder each job's bullets by JD-relevance: highest-relevance bullet first.
- Rewording should be BOLDER than "synonym swap". If JD language fits and the fact supports it:
  - Shift the verb ("built" → "shipped", "designed" → "architected", "developed" → "engineered") when it changes emphasis.
  - Promote the JD-relevant noun earlier in the sentence.
  - Collapse weak connectives; lead with the outcome.
- You MAY NOT: add a new technology, change a metric, swap the employer, or claim a different scope.

SKILLS:
- Reorder categories so the most JD-relevant ones come first.
- Within a category, reorder items so the JD-relevant ones lead.
- You MAY slightly rename a category label to match common vocabulary (e.g. "Cloud / DevOps" → "Cloud & Infrastructure") if it aids scannability.
- You MAY ADD new items to existing categories when the JD names a tech the candidate doesn't currently list. Place added items at the END of the most appropriate category, after the original items. Cap added items at ~3 per category and ~6 total across the resume — only add things the JD explicitly emphasizes.
- Added items must be plausibly adjacent to what the candidate already has (e.g. add "GraphQL" if they already list "REST APIs"; add "Kotlin" if they already list "Java"). Do NOT add a totally unrelated tech (e.g. don't add "Solidity" to a backend Python candidate).
- Do NOT add new categories. Use existing ones.

OUTPUT FORMAT — emit ONLY the XML below, nothing else. No prose, no markdown, no JSON, no explanation. Include a section only if the corresponding section was present in the input.

<summary>Your rewritten summary here.</summary>           ← only if the input had a <summary>

<job index="0">
  <b>First bullet for job 0 (highest JD-relevance)</b>
  <b>Second bullet for job 0</b>
  ... (exactly the same number of <b> tags as the input had for this job) ...
</job>
<job index="1">
  ...
</job>
... (one <job> block per input job, in input order by index) ...

<skill><category>Category name</category><items>item1, item2, item3</items></skill>
... (one <skill> block per input category, in your chosen order) ...

XML RULES:
- Straight ASCII double-quotes for attributes.
- Do NOT HTML-escape text inside <b>, <category>, <items>, <summary> — emit the final text as it should appear in the resume.
- <job index="N"> MUST match the input job_index exactly.
- No nested tags inside <b>, <category>, <items>, or <summary>."""



_TAILOR_CACHE_DIR = config.DATA_DIR / "tailor_cache"
_tailor_mem_cache: dict[str, dict] = {}

_SKILLS_CACHE_DIR = config.DATA_DIR / "skills_cache"
_skills_mem_cache: dict[str, list] = {}


def tailor_resume(
    resume: ResumeStruct, job: dict, *, system_prompt: Optional[str] = None
) -> dict:
    """Tailor `resume` for `job`. If `system_prompt` is non-empty, override
    the global default. Cached on disk by hash of (resume, job, system_prompt).
    """
    client = _client()
    sys_prompt = (system_prompt or "").strip() or TAILOR_SYSTEM

    input_parts: list[str] = []
    if resume.has_summary:
        input_parts += [f"<summary>{_xml_escape(resume.summary.strip())}</summary>", ""]
    for i, j in enumerate(resume.jobs):
        input_parts.append(
            f'<job index="{i}" title="{_xml_escape(j.title_text.strip())}" '
            f'bullet_count="{len(j.bullets)}">'
        )
        for b in j.bullets:
            input_parts.append(f"  <b>{_xml_escape(b.strip())}</b>")
        input_parts.append("</job>")
    if resume.jobs:
        input_parts.append("")
    for s in resume.skills:
        input_parts.append(
            f"<skill><category>{_xml_escape(s.category.strip())}</category>"
            f"<items>{_xml_escape(s.items.strip())}</items></skill>"
        )
    resume_xml = "\n".join(input_parts)

    cache_blob = json.dumps({
        "model": config.TAILOR_MODEL,
        "system": sys_prompt,
        "resume_xml": resume_xml,
        "company": job.get("company", ""),
        "title": job.get("title", ""),
        "location": job.get("location", ""),
        "description": (job.get("description") or "")[:15000],
    }, ensure_ascii=False, sort_keys=True)
    cache_key = hashlib.sha256(cache_blob.encode("utf-8")).hexdigest()
    cached = _tailor_mem_cache.get(cache_key)
    if cached is None:
        cache_path = _TAILOR_CACHE_DIR / f"{cache_key}.json"
        if cache_path.exists():
            try:
                cached = json.loads(cache_path.read_text(encoding="utf-8"))
                _tailor_mem_cache[cache_key] = cached
            except (OSError, json.JSONDecodeError):
                cached = None
    if cached is not None:
        return cached

    user_msg = [
        {
            "type": "text",
            "text": "Candidate resume:\n<resume>\n" + resume_xml + "\n</resume>",
            "cache_control": {"type": "ephemeral"},
        },
        {
            "type": "text",
            "text": (
                "Target job posting:\n<jd>\n"
                f"Company: {job.get('company','')}\n"
                f"Title: {job.get('title','')}\n"
                f"Location: {job.get('location','')}\n\n"
                f"Description:\n{(job.get('description') or '')[:15000]}\n"
                "</jd>\n\n"
                "Before writing the output, silently work through these steps "
                "(do NOT include them in the output):\n"
                "  1) List the JD's top 5-7 must-have requirements in priority order.\n"
                "  2) For each, decide if the resume has an exact match, an adjacent match, or no support.\n"
                "  3) Decide the summary's domain word from the JD.\n"
                "  4) Rank each job's bullets by relevance before writing.\n"
                "Then emit ONLY the XML format from the system prompt. "
                "Be bold in rewording while staying 100% truthful to the input facts."
            ),
        },
    ]

    resp = client.messages.create(
        model=config.TAILOR_MODEL,
        max_tokens=4000,
        temperature=0,
        system=sys_prompt,
        messages=[{"role": "user", "content": user_msg}],
    )
    text = "".join(
        b.text for b in resp.content if getattr(b, "type", None) == "text"
    )
    out = _parse_xml_output(text, resume)
    _repair_output(out, resume)
    _validate_tailored(out, resume)

    _tailor_mem_cache[cache_key] = out
    try:
        _TAILOR_CACHE_DIR.mkdir(parents=True, exist_ok=True)
        (_TAILOR_CACHE_DIR / f"{cache_key}.json").write_text(
            json.dumps(out, ensure_ascii=False, indent=2), encoding="utf-8"
        )
    except OSError:
        pass

    return out


SKILLS_SYSTEM = """You rewrite the SKILLS section of a resume to align with a target job posting, for ATS keyword matching. Be aggressive about surfacing relevant skills — but only ones the candidate actually has.

You are given:
- The candidate's current skills as numbered rows, each "Category: item, item, ...".
- The JD's prioritized skills (most important first).
- An ALLOWED list — the ONLY skills you may ADD (the candidate has exact or adjacent/transferable evidence for these).

RULES — any violation invalidates the output:
1. Output EXACTLY the same number of rows as the input.
2. You may ADD a skill only if it is in the ALLOWED list. NEVER add a skill that is not in the ALLOWED list, even if the JD wants it.
3. Never remove a skill the candidate already listed. No duplicates within a row.
4. Place each added skill in the most fitting row.
5. Reorder the rows so the most JD-relevant categories come first; within each row, lead with the most JD-relevant items.
6. Canonicalize names to the JD's spelling (write "Kubernetes" if the JD says Kubernetes and the resume said "k8s").
7. You MAY rename a category label to clearer / JD-aligned wording, as long as it stays truthful to the row's contents.

OUTPUT — emit STRICT JSON inside <json>...</json> tags, nothing else:
{"rows": [{"category": "Languages", "items": "Python, Go, ..."}, ...]}
Exactly the same number of rows as the input, in your chosen order."""


def tailor_skills(
    skill_rows: list[tuple[str, str]],
    jd_terms: list[str],
    allowed_terms: list[str],
) -> Optional[list[tuple[str, str]]]:
    """LLM-rewrite the skills section: reorder + canonicalize + add allowed JD
    skills into the best-fitting existing rows. Returns the rewritten rows (same
    count) or None on failure / invalid output (caller falls back)."""
    rows = [(str(c or "").strip(), str(i or "").strip()) for c, i in skill_rows]
    if not rows:
        return None
    n = len(rows)
    rows_txt = "\n".join(f"{idx + 1}. {c}: {it}" for idx, (c, it) in enumerate(rows))
    user = (
        f"CURRENT SKILLS ({n} rows):\n{rows_txt}\n\n"
        f"JD PRIORITY SKILLS:\n{', '.join(jd_terms[:50])}\n\n"
        f"ALLOWED TO ADD (only these):\n{', '.join(allowed_terms[:80]) or '(none)'}\n\n"
        f"Return JSON with exactly {n} rows."
    )

    # Disk + mem cache keyed on the full input. Inputs are deterministic for a
    # given (profile, JD), so a repeated tailoring of the same pair is free.
    cache_blob = json.dumps(
        {"model": config.JUDGE_MODEL, "system": SKILLS_SYSTEM, "user": user},
        ensure_ascii=False, sort_keys=True,
    )
    cache_key = hashlib.sha256(cache_blob.encode("utf-8")).hexdigest()
    cached = _skills_mem_cache.get(cache_key)
    if cached is None:
        cache_path = _SKILLS_CACHE_DIR / f"{cache_key}.json"
        if cache_path.exists():
            try:
                cached = json.loads(cache_path.read_text(encoding="utf-8"))
                _skills_mem_cache[cache_key] = cached
            except (OSError, json.JSONDecodeError):
                cached = None
    if cached is not None:
        # Stored as list[[cat, items]]; return as list[tuple].
        return [tuple(r) for r in cached] if cached else None

    try:
        resp = _client().messages.create(
            model=config.JUDGE_MODEL,
            max_tokens=1500,
            temperature=0,
            system=SKILLS_SYSTEM,
            messages=[{"role": "user", "content": user}],
        )
        txt = "".join(b.text for b in resp.content if getattr(b, "type", None) == "text")
        data = _extract_tagged_json(txt)
        out_rows = data.get("rows") if isinstance(data, dict) else None
        if not isinstance(out_rows, list):
            return None
        out = [
            (str(r.get("category", "")).strip(), str(r.get("items", "")).strip())
            for r in out_rows if isinstance(r, dict)
        ]
        out = [r for r in out if r[0] or r[1]]
        # Must preserve the row count so apply_tailoring maps content to the same
        # docx paragraph slots.
        if len(out) != n:
            return None
        # Persist (best-effort) so repeated runs of this (profile, JD) skip the call.
        try:
            _SKILLS_CACHE_DIR.mkdir(parents=True, exist_ok=True)
            _SKILLS_CACHE_DIR.joinpath(f"{cache_key}.json").write_text(
                json.dumps([list(r) for r in out], ensure_ascii=False),
                encoding="utf-8",
            )
        except OSError:
            pass
        _skills_mem_cache[cache_key] = [list(r) for r in out]
        return out
    except Exception:
        return None


def _xml_escape(s: str) -> str:
    return (
        s.replace("&", "&amp;")
         .replace("<", "&lt;")
         .replace(">", "&gt;")
         .replace('"', "&quot;")
    )


def _xml_unescape(s: str) -> str:
    return (
        s.replace("&quot;", '"')
         .replace("&lt;", "<")
         .replace("&gt;", ">")
         .replace("&amp;", "&")
    )


def _parse_xml_output(text: str, resume: ResumeStruct) -> dict:
    """Parse the model's XML response into the canonical dict."""
    summary_m = re.search(r"<summary>(.*?)</summary>", text, re.DOTALL)
    summary = _xml_unescape(summary_m.group(1).strip()) if summary_m else ""

    jobs_by_index: dict = {}
    for m in re.finditer(
        r'<job[^>]*\bindex="(\d+)"[^>]*>(.*?)</job>', text, re.DOTALL
    ):
        idx = int(m.group(1))
        body = m.group(2)
        bullets = [
            _xml_unescape(b.strip())
            for b in re.findall(r"<b>(.*?)</b>", body, re.DOTALL)
        ]
        jobs_by_index[idx] = bullets
    bullets_per_job = [jobs_by_index.get(i, []) for i in range(len(resume.jobs))]

    skill_cats: list = []
    skill_items: list = []
    for m in re.finditer(
        r"<skill>\s*<category>(.*?)</category>\s*<items>(.*?)</items>\s*</skill>",
        text,
        re.DOTALL,
    ):
        skill_cats.append(_xml_unescape(m.group(1).strip()))
        skill_items.append(_xml_unescape(m.group(2).strip()))

    return {
        "summary": summary,
        "bullets": bullets_per_job,
        "skill_categories": skill_cats,
        "skill_items": skill_items,
    }


def _repair_output(out: dict, resume: ResumeStruct) -> None:
    """Fill in anything the model dropped by falling back to the original resume.

    Claude occasionally omits a skill category or the summary block. Rather
    than fail the whole row (which wastes the tailoring call we just paid for),
    we patch the missing pieces with the original content. The model's ordering
    and added items are preserved; we only add back what it left out.
    """
    # Summary: if the input had one but the model didn't emit one, keep the
    # original summary unchanged rather than erroring out.
    if resume.has_summary:
        s = out.get("summary")
        if not isinstance(s, str) or not s.strip():
            out["summary"] = resume.summary

    # Bullets: per-job padding. If a job's bullet list is too short, append the
    # missing original bullets (at the end) so the count matches the layout.
    bullets = out.get("bullets")
    if not isinstance(bullets, list):
        bullets = []
    while len(bullets) < len(resume.jobs):
        rj = resume.jobs[len(bullets)]
        bullets.append([b.strip() for b in rj.bullets])
    for i, (inner, rj) in enumerate(zip(bullets, resume.jobs)):
        if not isinstance(inner, list):
            bullets[i] = [b.strip() for b in rj.bullets]
            continue
        if len(inner) < len(rj.bullets):
            # Missing bullets: fill with originals that the model didn't keep,
            # matched case-insensitively so we don't duplicate what's already there.
            have = {str(b).strip().lower() for b in inner}
            for orig in rj.bullets:
                if orig.strip().lower() not in have:
                    inner.append(orig.strip())
                if len(inner) >= len(rj.bullets):
                    break
            # If still short (the model reworded everything), pad at the end.
            while len(inner) < len(rj.bullets):
                inner.append(rj.bullets[len(inner)].strip())
        elif len(inner) > len(rj.bullets):
            # Model added bullets — trim to the original count (layout fixed).
            bullets[i] = inner[: len(rj.bullets)]
    out["bullets"] = bullets

    # Skill categories + items: if short, append the originals the model
    # didn't include (matched case-insensitively by category label).
    if resume.skills:
        sc = list(out.get("skill_categories") or [])
        si = list(out.get("skill_items") or [])
        # Trim to len(resume.skills) so the model can't overshoot the layout —
        # also keeps sc/si the same length when one is longer than the other.
        n = min(len(sc), len(si), len(resume.skills))
        sc, si = sc[:n], si[:n]
        if len(sc) < len(resume.skills):
            have = {str(c).strip().lower() for c in sc}
            for s in resume.skills:
                if s.category.strip().lower() not in have:
                    sc.append(s.category.strip())
                    si.append(s.items.strip())
                if len(sc) >= len(resume.skills):
                    break
            # Last resort pad (e.g. model renamed every category).
            while len(sc) < len(resume.skills):
                k = len(sc)
                sc.append(resume.skills[k].category.strip())
                si.append(resume.skills[k].items.strip())
        out["skill_categories"] = sc
        out["skill_items"] = si


def _validate_tailored(out: dict, resume: ResumeStruct) -> None:
    """Hard-check after repair. Any remaining mismatch is a real bug."""
    if resume.has_summary:
        if not isinstance(out.get("summary"), str) or not out["summary"].strip():
            raise ValueError("Tailored output missing summary even after repair.")

    bullets = out.get("bullets")
    if not isinstance(bullets, list) or len(bullets) != len(resume.jobs):
        raise ValueError(
            f"'bullets' must be a list of {len(resume.jobs)} inner lists."
        )
    for i, (inner, rj) in enumerate(zip(bullets, resume.jobs)):
        if not isinstance(inner, list) or len(inner) != len(rj.bullets):
            raise ValueError(
                f"bullets[{i}] must have {len(rj.bullets)} strings."
            )

    if resume.skills:
        sc = out.get("skill_categories")
        si = out.get("skill_items")
        if not isinstance(sc, list) or len(sc) != len(resume.skills):
            raise ValueError(
                f"'skill_categories' must be a list of {len(resume.skills)} strings."
            )
        if not isinstance(si, list) or len(si) != len(resume.skills):
            raise ValueError(
                f"'skill_items' must be a list of {len(resume.skills)} strings."
            )


def apply_tailoring(
    src_docx: Path, resume: ResumeStruct, tailored: dict, dst_docx: Path
) -> None:
    doc = Document(str(src_docx))
    paras = doc.paragraphs

    if resume.has_summary and tailored.get("summary", "").strip():
        set_paragraph_text(paras[resume.summary_idx], tailored["summary"].strip())

    for job, inner_bullets in zip(resume.jobs, tailored.get("bullets", [])):
        for idx, new_bullet in zip(job.bullet_indices, inner_bullets):
            set_paragraph_text(paras[idx], new_bullet.strip())

    if resume.skills:
        for skill, cat, items in zip(
            resume.skills,
            tailored.get("skill_categories", []),
            tailored.get("skill_items", []),
        ):
            if skill.category_idx == skill.items_idx:
                # Category + items share one paragraph — write them together
                # (bold label + plain items) so neither clobbers the other.
                _apply_skill_row(paras[skill.category_idx], cat, items)
            else:
                set_paragraph_text(paras[skill.category_idx], cat.strip())
                set_paragraph_text(paras[skill.items_idx], items.strip())

    dst_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst_docx))
